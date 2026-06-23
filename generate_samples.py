"""
Generates sample financial documents for pipeline testing.
Outputs 2 PDFs and 2 Word docs into data/sample_docs/.
"""

import os
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("data/sample_docs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# PDF 1 — Client Portfolio Statement
# ---------------------------------------------------------------------------

def create_portfolio_statement():
    path = OUTPUT_DIR / "client_portfolio_statement.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    heading = ParagraphStyle("heading", parent=styles["Heading1"], spaceAfter=6)
    subheading = ParagraphStyle("subheading", parent=styles["Heading2"], spaceAfter=4)
    normal = styles["Normal"]

    story = []

    story.append(Paragraph("QUARTERLY PORTFOLIO STATEMENT", heading))
    story.append(Paragraph("Client: James & Sarah Johnson", normal))
    story.append(Paragraph("Account Number: AUS-2024-00471", normal))
    story.append(Paragraph("Statement Period: 1 January 2024 – 31 March 2024", normal))
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Portfolio Summary", subheading))
    summary_data = [
        ["", ""],
        ["Total Portfolio Value", "$842,350.00"],
        ["Period Return", "+4.2%"],
        ["1-Year Return", "+11.8%"],
        ["Benchmark (ASX 200)", "+9.6%"],
    ]
    summary_table = Table(summary_data, colWidths=[3 * inch, 2.5 * inch])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Asset Allocation", subheading))
    allocation_data = [
        ["Asset Class", "Value", "Weight", "Target Weight"],
        ["Australian Equities", "$336,940", "40.0%", "40%"],
        ["International Equities", "$210,587", "25.0%", "25%"],
        ["Fixed Income", "$168,470", "20.0%", "20%"],
        ["Property (REITs)", "$84,235", "10.0%", "10%"],
        ["Cash & Equivalents", "$42,118", "5.0%", "5%"],
        ["Total", "$842,350", "100%", "100%"],
    ]
    alloc_table = Table(allocation_data, colWidths=[2.5 * inch, 1.5 * inch, 1.2 * inch, 1.5 * inch])
    alloc_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -2), "Helvetica"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.whitesmoke, colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(alloc_table)
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Top Holdings", subheading))
    holdings_data = [
        ["Security", "Ticker", "Shares", "Price", "Value", "Weight"],
        ["Commonwealth Bank", "CBA.AX", "420", "$118.40", "$49,728", "5.9%"],
        ["BHP Group", "BHP.AX", "890", "$44.20", "$39,338", "4.7%"],
        ["CSL Limited", "CSL.AX", "185", "$285.60", "$52,836", "6.3%"],
        ["Westpac Banking", "WBC.AX", "1,200", "$26.80", "$32,160", "3.8%"],
        ["Macquarie Group", "MQG.AX", "210", "$192.50", "$40,425", "4.8%"],
    ]
    holdings_table = Table(holdings_data,
                           colWidths=[1.8 * inch, 0.9 * inch, 0.7 * inch, 0.8 * inch, 1 * inch, 0.8 * inch])
    holdings_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(holdings_table)
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Adviser Notes", subheading))
    notes = (
        "The Johnson portfolio continues to perform above benchmark, driven primarily by strong returns "
        "in the Australian equities sleeve. CSL and CBA were the top contributors this quarter. "
        "Fixed income allocation has been maintained at target weight given the current interest rate "
        "environment. We recommend reviewing the property REIT exposure at the next client meeting "
        "given recent softness in commercial real estate valuations. Overall risk profile remains "
        "Balanced-Growth, consistent with the client's stated objectives and 12-year investment horizon."
    )
    story.append(Paragraph(notes, normal))

    doc.build(story)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# PDF 2 — Investment Proposal
# ---------------------------------------------------------------------------

def create_investment_proposal():
    path = OUTPUT_DIR / "investment_proposal.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    heading = ParagraphStyle("heading", parent=styles["Heading1"], spaceAfter=6)
    subheading = ParagraphStyle("subheading", parent=styles["Heading2"], spaceAfter=4)
    normal = styles["Normal"]

    story = []

    story.append(Paragraph("INVESTMENT PROPOSAL", heading))
    story.append(Paragraph("Prepared for: Michael Chen", normal))
    story.append(Paragraph("Prepared by: Pinnacle Wealth Advisory", normal))
    story.append(Paragraph("Date: 15 March 2024", normal))
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Executive Summary", subheading))
    summary = (
        "This proposal outlines a recommended investment strategy for Michael Chen, aged 42, "
        "with an investable sum of $250,000. Based on our risk assessment, Michael's profile is "
        "Growth-oriented with a 15-year investment horizon targeting retirement at age 57. "
        "The recommended portfolio targets an annualised return of 8–10% with a medium-high "
        "risk tolerance."
    )
    story.append(Paragraph(summary, normal))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Client Goals", subheading))
    goals_data = [
        ["Goal", "Target Amount", "Timeframe", "Priority"],
        ["Retirement income", "$2,000,000", "15 years", "High"],
        ["Children's education", "$120,000", "8 years", "High"],
        ["Holiday home", "$600,000", "10 years", "Medium"],
        ["Emergency buffer", "$50,000", "2 years", "High"],
    ]
    goals_table = Table(goals_data, colWidths=[2 * inch, 1.5 * inch, 1.2 * inch, 1.2 * inch])
    goals_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a5276")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(goals_table)
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Recommended Portfolio", subheading))
    portfolio_data = [
        ["Asset Class", "Allocation", "Expected Return", "Expected Risk"],
        ["Global Equities (Unhedged)", "45%", "9.5% p.a.", "High"],
        ["Australian Equities", "25%", "8.8% p.a.", "Medium-High"],
        ["Emerging Markets", "10%", "11.2% p.a.", "High"],
        ["Infrastructure", "10%", "6.5% p.a.", "Medium"],
        ["Cash & Short-Term Bonds", "10%", "4.8% p.a.", "Low"],
    ]
    portfolio_table = Table(portfolio_data, colWidths=[2.2 * inch, 1.2 * inch, 1.5 * inch, 1.5 * inch])
    portfolio_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a5276")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(portfolio_table)
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Fee Disclosure", subheading))
    fees = (
        "Ongoing advice fee: 0.75% p.a. of portfolio value (capped at $5,000 p.a.). "
        "Investment management fees (MER): 0.18%–0.45% p.a. depending on fund selection. "
        "Estimated total cost of advice and investment: approximately 1.0%–1.2% p.a. "
        "No entry or exit fees apply. Performance fees do not apply to the recommended product set."
    )
    story.append(Paragraph(fees, normal))

    doc.build(story)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# Word Doc 1 — Financial Plan
# ---------------------------------------------------------------------------

def create_financial_plan():
    path = OUTPUT_DIR / "financial_plan.docx"
    doc = Document()

    title = doc.add_heading("PERSONAL FINANCIAL PLAN", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Client: Emily & Robert Patel")
    doc.add_paragraph("Adviser: Pinnacle Wealth Advisory")
    doc.add_paragraph("Plan Date: 20 February 2024")
    doc.add_paragraph("Review Date: 20 February 2025")
    doc.add_paragraph()

    doc.add_heading("1. Client Profile", level=1)
    profile = doc.add_paragraph()
    profile.add_run("Emily Patel").bold = True
    doc.add_paragraph("Age: 38 | Occupation: Senior Software Engineer | Income: $145,000 p.a.")
    profile2 = doc.add_paragraph()
    profile2.add_run("Robert Patel").bold = True
    doc.add_paragraph("Age: 41 | Occupation: GP (General Practitioner) | Income: $220,000 p.a.")
    doc.add_paragraph("Combined household income: $365,000 p.a.")
    doc.add_paragraph("Dependants: Two children aged 6 and 9")
    doc.add_paragraph()

    doc.add_heading("2. Goals and Objectives", level=1)
    goals = [
        ("Retirement", "Both to retire by age 60 with combined income of $120,000 p.a. in today's dollars."),
        ("Education", "Fund private secondary and tertiary education for both children (~$80,000 each)."),
        ("Debt Reduction", "Eliminate remaining mortgage ($380,000) within 8 years."),
        ("Insurance", "Ensure adequate life, TPD, and income protection coverage for both partners."),
        ("Estate Planning", "Update wills and establish a testamentary trust for the children."),
    ]
    for title_text, detail in goals:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(detail)
    doc.add_paragraph()

    doc.add_heading("3. Current Financial Position", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Asset / Liability"
    hdr[1].text = "Value"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    rows = [
        ("Family home (Sydney)", "$1,650,000"),
        ("Mortgage outstanding", "-$380,000"),
        ("Emily — Superannuation", "$185,000"),
        ("Robert — Superannuation", "$320,000"),
        ("Share portfolio (joint)", "$95,000"),
        ("Savings account", "$42,000"),
        ("Robert's practice equity (est.)", "$480,000"),
        ("Net Worth (estimated)", "$2,392,000"),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    doc.add_paragraph()

    doc.add_heading("4. Recommendations", level=1)
    recs = [
        "Increase voluntary superannuation contributions for both Emily and Robert to maximise concessional cap ($27,500 p.a. each).",
        "Establish an investment bond for children's education funding — target $500/month per child.",
        "Implement a debt recycling strategy to accelerate mortgage reduction while building tax-effective investments.",
        "Review and increase Robert's income protection insurance — current policy only covers 75% of income to age 55; recommend extending to age 65.",
        "Engage estate planning solicitor to draft testamentary trust wills.",
        "Review asset allocation of joint share portfolio — currently overweight technology sector (42%).",
    ]
    for rec in recs:
        doc.add_paragraph(rec, style="List Number")
    doc.add_paragraph()

    doc.add_heading("5. Risk Profile", level=1)
    doc.add_paragraph(
        "Based on the risk profiling questionnaire completed on 15 February 2024, both clients have been "
        "assessed as Balanced–Growth investors. They are comfortable with short-term fluctuations in portfolio "
        "value and have a long-term investment horizon of 20+ years for retirement assets. A moderate drawdown "
        "of up to 20% in any 12-month period is considered acceptable."
    )
    doc.add_paragraph()

    doc.add_heading("6. Next Steps", level=1)
    next_steps = [
        "Sign and return the Statement of Advice (SOA) by 5 March 2024.",
        "Provide superannuation fund details for contribution redirection.",
        "Book appointment with recommended estate planning solicitor.",
        "Complete insurance needs analysis — scheduled for 28 February 2024.",
    ]
    for step in next_steps:
        doc.add_paragraph(step, style="List Bullet")

    doc.save(str(path))
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# Word Doc 2 — Risk Assessment
# ---------------------------------------------------------------------------

def create_risk_assessment():
    path = OUTPUT_DIR / "risk_assessment.docx"
    doc = Document()

    title = doc.add_heading("CLIENT RISK ASSESSMENT QUESTIONNAIRE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Client: Thomas Nguyen")
    doc.add_paragraph("Date Completed: 8 March 2024")
    doc.add_paragraph("Adviser: Pinnacle Wealth Advisory")
    doc.add_paragraph()

    doc.add_heading("Section 1 — Investment Experience", level=1)
    q_and_a = [
        ("How long have you been investing?", "More than 10 years"),
        ("Which asset classes have you invested in?", "Shares, property, managed funds, ETFs"),
        ("How familiar are you with investment risk?", "Very familiar — I understand that returns fluctuate and capital can be lost"),
    ]
    for q, a in q_and_a:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        doc.add_paragraph(f"Response: {a}")
        doc.add_paragraph()

    doc.add_heading("Section 2 — Risk Tolerance", level=1)
    risk_qa = [
        ("If your portfolio fell by 20% in a year, what would you do?",
         "Hold — I understand markets recover and would not sell"),
        ("What is your primary investment objective?",
         "Long-term capital growth, accepting higher short-term volatility"),
        ("How important is it that your investments keep pace with inflation?",
         "Very important — I want real returns above inflation over time"),
    ]
    for q, a in risk_qa:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        doc.add_paragraph(f"Response: {a}")
        doc.add_paragraph()

    doc.add_heading("Section 3 — Financial Situation", level=1)
    fin_qa = [
        ("Annual income", "$185,000"),
        ("Investment timeframe", "15–20 years"),
        ("Liquidity needs", "Low — I have a separate emergency fund of $60,000"),
        ("Major financial obligations", "Mortgage ($220,000 remaining), no other debt"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Response"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for q, a in fin_qa:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = a
    doc.add_paragraph()

    doc.add_heading("Section 4 — Risk Profile Result", level=1)
    result = doc.add_paragraph()
    result.add_run("Assessed Risk Profile: ").bold = True
    result.add_run("Growth")
    doc.add_paragraph(
        "Thomas Nguyen has been assessed as a Growth investor. This profile is suitable for investors "
        "who seek higher long-term returns and are willing to accept significant short-term volatility. "
        "The recommended asset allocation is approximately 75% growth assets (equities, property) and "
        "25% defensive assets (bonds, cash). A maximum drawdown tolerance of 30% over a 12-month "
        "period has been noted."
    )
    doc.add_paragraph()

    doc.add_heading("Section 5 — Adviser Confirmation", level=1)
    doc.add_paragraph(
        "I confirm that this risk assessment was completed in the presence of the client and accurately "
        "reflects their stated investment objectives, risk tolerance, and financial circumstances as at "
        "the date of this assessment."
    )
    doc.add_paragraph()
    doc.add_paragraph("Adviser signature: _______________________")
    doc.add_paragraph("Date: ___________________________________")

    doc.save(str(path))
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating sample financial documents...\n")
    create_portfolio_statement()
    create_investment_proposal()
    create_financial_plan()
    create_risk_assessment()
    print("\nDone. Files saved to data/sample_docs/")
