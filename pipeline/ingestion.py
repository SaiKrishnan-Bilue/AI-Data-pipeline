"""
Step 2 — Document Ingestion

Reads PDF and Word documents, extracts raw text AND structured tables,
and saves each result as a JSON file in data/processed/.
"""

import json
from pathlib import Path
from typing import Optional

import fitz  # pymupdf
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def _parse_table_rows(raw_rows: list) -> Optional[dict]:
    """
    Converts a raw list-of-lists into {"headers": [...], "rows": [...]}.
    Returns None if the table is empty or has fewer than 2 rows.
    """
    # Strip None and whitespace from every cell
    cleaned = [
        [str(cell).strip() if cell is not None else "" for cell in row]
        for row in raw_rows
    ]
    # Drop completely empty rows
    cleaned = [row for row in cleaned if any(cell for cell in row)]

    if len(cleaned) < 2:
        return None

    return {"headers": cleaned[0], "rows": cleaned[1:]}


def extract_from_pdf(file_path: Path) -> list:
    """
    Returns one dict per page with:
      - page: page number
      - text: raw text with table content removed
      - tables: list of structured tables found on that page
    """
    pages = []
    with fitz.open(str(file_path)) as pdf:
        for i, page in enumerate(pdf, start=1):
            # Find tables first
            tables = []
            tab_finder = page.find_tables()
            for tab in tab_finder.tables:
                parsed = _parse_table_rows(tab.extract())
                if parsed:
                    tables.append(parsed)

            # Extract text — pymupdf still includes table text in get_text(),
            # so we use "blocks" mode and skip blocks that overlap table bboxes
            table_bboxes = [fitz.Rect(tab.bbox) for tab in tab_finder.tables]
            text_blocks = []
            for block in page.get_text("blocks"):
                block_rect = fitz.Rect(block[:4])
                overlaps_table = any(block_rect.intersects(tb) for tb in table_bboxes)
                if not overlaps_table and block[4].strip():
                    text_blocks.append(block[4].strip())

            pages.append({
                "page": i,
                "text": "\n".join(text_blocks),
                "tables": tables,
            })

    return pages


def extract_from_docx(file_path: Path) -> list:
    """
    Word docs don't have pages, so we return a single entry.
    Paragraphs and tables are extracted separately.
    Tables in Word are not part of doc.paragraphs — they live in doc.tables.
    """
    doc = Document(str(file_path))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        raw_rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        parsed = _parse_table_rows(raw_rows)
        if parsed:
            tables.append(parsed)

    return [{
        "page": 1,
        "text": "\n".join(paragraphs),
        "tables": tables,
    }]


def ingest_document(file_path: Path) -> dict:
    """
    Detects file type, extracts text and tables, returns a structured result.
    """
    ext = file_path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        pages = extract_from_pdf(file_path)
    elif ext == ".docx":
        pages = extract_from_docx(file_path)

    total_chars = sum(len(p["text"]) for p in pages)
    total_tables = sum(len(p["tables"]) for p in pages)

    return {
        "file_name": file_path.name,
        "file_type": ext,
        "page_count": len(pages),
        "total_characters": total_chars,
        "total_tables": total_tables,
        "pages": pages,
    }


def run_ingestion(input_dir: Path, output_dir: Path) -> None:
    """
    Processes all supported documents in input_dir.
    Saves each result as a JSON file in output_dir.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    files = [f for f in input_dir.iterdir() if f.suffix.lower() in SUPPORTED_EXTENSIONS]

    if not files:
        print(f"No supported documents found in {input_dir}")
        return

    for file_path in sorted(files):
        print(f"Ingesting: {file_path.name} ...", end=" ")
        result = ingest_document(file_path)

        output_file = output_dir / (file_path.stem + "_ingested.json")
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        print(f"done  ({result['page_count']} page(s), {result['total_characters']:,} chars, {result['total_tables']} table(s))")

    print(f"\nAll files saved to {output_dir}")


if __name__ == "__main__":
    run_ingestion(
        input_dir=Path("data/sample_docs"),
        output_dir=Path("data/processed"),
    )
